#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import sys
import os
import markdown
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from bs4 import BeautifulSoup  # 用于解析 HTML

def md_to_docx(md_file, docx_file=None):
    if not os.path.exists(md_file):
        print(f"❌ 文件不存在: {md_file}")
        return False

    if docx_file is None:
        docx_file = os.path.splitext(md_file)[0] + ".docx"

    # 读取 Markdown 文件
    with open(md_file, "r", encoding="utf-8") as f:
        md_text = f.read()

    # 转为 HTML
    html = markdown.markdown(
        md_text,
        extensions=[
            'tables',       # 支持表格（需额外处理）
            'fenced_code',  # 支持 ``` 代码块
            'nl2br'         # 换行转 <br>
        ]
    )

    # 创建 Word 文档
    doc = Document()
    soup = BeautifulSoup(html, "html.parser")

    def add_paragraph_with_format(element):
        text = element.get_text()
        p = doc.add_paragraph()

        # 处理内联样式（加粗、斜体等）
        for child in element.children:
            if child.name == 'strong' or child.name == 'b':
                run = p.add_run(child.get_text())
                run.bold = True
            elif child.name == 'em' or child.name == 'i':
                run = p.add_run(child.get_text())
                run.italic = True
            elif child.name == 'code':
                run = p.add_run(child.get_text())
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
            elif child.name is None:  # 纯文本
                p.add_run(str(child))
            else:
                p.add_run(child.get_text())

    # 遍历 HTML 元素
    for element in soup.children:
        if element.name == 'h1':
            doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            doc.add_heading(element.get_text(), level=3)
        elif element.name == 'h4':
            doc.add_heading(element.get_text(), level=4)
        elif element.name == 'p':
            add_paragraph_with_format(element)
        elif element.name == 'ul':
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(li.get_text())
        elif element.name == 'ol':
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Number')
                p.add_run(li.get_text())
        elif element.name == 'pre':
            code = element.get_text()
            p = doc.add_paragraph()
            run = p.add_run(code)
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        elif element.name == 'table':
            # 简单表格支持（需安装 lxml 或 html5lib）
            try:
                table = doc.add_table(rows=0, cols=len(element.find('tr').find_all(['td', 'th'])))
                table.style = 'Table Grid'
                for row in element.find_all('tr'):
                    cells = table.add_row().cells
                    for i, cell in enumerate(row.find_all(['td', 'th'])):
                        cells[i].text = cell.get_text()
            except Exception as e:
                doc.add_paragraph(f"[表格转换失败: {e}]")
        elif element.name is None:
            continue  # 忽略空白
        else:
            # 兜底：直接添加文本
            doc.add_paragraph(element.get_text())

    doc.save(docx_file)
    print(f"✅ 转换成功！\n   源文件: {md_file}\n   目标文件: {docx_file}")
    return True

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python md_to_docx.py <输入.md> [输出.docx]")
        sys.exit(1)

    input_md = sys.argv[1]
    output_docx = sys.argv[2] if len(sys.argv) > 2 else None
    md_to_docx(input_md, output_docx)